"""
Sinh file Final-Report-v2.docx — bản report updated cho dự án BĐS Multi-agent.

Chạy: `python scripts/build_report.py` (từ root project)
Output: `Final-Report-v2.docx` ở root project.

File cấu trúc 6 chương khớp report cũ (Final-Report Updated.pdf), nội dung đã update
theo code hiện tại: Planner Agent, Query Rewriter, OpenAI stack, frontend Streamlit
multi-chat, dashboard trim, ES isolate per chat, không còn MongoDB user auth.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# Windows console cp1252 chết với tiếng Việt nếu print
sys.stdout.reconfigure(encoding="utf-8")

OUTPUT = Path(__file__).resolve().parent.parent / "Final-Report-v2.docx"


# ============================================================
# Helpers — apply Vietnamese academic style
# ============================================================

def setup_styles(doc: Document) -> None:
    """Set default font + heading sizes giống report cũ (Times New Roman 13)."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    for name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 13, True),
        ("Heading 3", 13, True),
    ]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)


def page_setup(doc: Document) -> None:
    """A4 margins giống chuẩn báo cáo BKHN."""
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.bold = True


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.bold = True


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    run = p.add_run(text)
    run.font.bold = True
    run.font.italic = False


def para(doc: Document, text: str, indent_first: bool = True, justify: bool = True) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + 0.5 * level)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)


def bold_label(doc: Document, label: str, body: str) -> None:
    """In-line bold label + plain body, vd: '**Why:** ...'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(label)
    r.font.bold = True
    p.add_run(" " + body)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.bold = True
    r.font.italic = True


# ============================================================
# COVER PAGE
# ============================================================

def build_cover(doc: Document) -> None:
    def center(text: str, size: int = 14, bold: bool = True, italic: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = bold
        r.font.italic = italic
        r.font.size = Pt(size)

    center("ĐẠI HỌC BÁCH KHOA HÀ NỘI", 14)
    center("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG", 14)
    for _ in range(8):
        doc.add_paragraph()
    center("BÁO CÁO BÀI TẬP LỚN", 20)
    doc.add_paragraph()
    center("NHẬP MÔN KHOA HỌC DỮ LIỆU", 14)
    doc.add_paragraph()
    center("ĐỀ TÀI: HỆ THỐNG ĐA TÁC TỬ", 16)
    center("TƯ VẤN BẤT ĐỘNG SẢN", 16)
    for _ in range(3):
        doc.add_paragraph()
    center("Giáo viên hướng dẫn: PGS.TS. Phạm Văn Hải", 13)
    center("Mã lớp: 162301    –    Nhóm 22", 13)
    doc.add_paragraph()

    # Bảng danh sách sinh viên
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Phạm Việt Hoàng", "20224854"),
        ("Nguyễn Thị Trà My", "20225049"),
        ("Đoàn Mạnh Hùng", "20224995"),
        ("Phạm Thanh An", "20224911"),
    ]
    for i, (name, msv) in enumerate(rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = name
        c1.text = msv
        for c in (c0, c1):
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()
    center("Hà Nội, Tháng 5 năm 2026", 13, bold=False, italic=True)
    page_break(doc)


# ============================================================
# LỜI NÓI ĐẦU
# ============================================================

def build_foreword(doc: Document) -> None:
    add_h1(doc, "LỜI NÓI ĐẦU")
    doc.add_paragraph()

    para(doc,
        "Cuộc cách mạng công nghiệp 4.0 đã thúc đẩy sự phát triển vượt bậc của công nghệ thông "
        "tin, dẫn đến sự bùng nổ của dữ liệu lớn trong nhiều lĩnh vực, bao gồm bất động sản. "
        "Các bài đăng bán bất động sản trên các nền tảng trực tuyến tạo ra một khối lượng dữ "
        "liệu khổng lồ, đa dạng và phức tạp, chứa đựng thông tin giá trị về thị trường.")

    para(doc,
        "Tại Việt Nam, thị trường bất động sản là một trong những lĩnh vực có tốc độ biến động "
        "nhanh và dữ liệu cực kỳ phân tán. Với hàng nghìn tin đăng mỗi ngày trên các nền tảng "
        "khác nhau, người dùng và các nhà đầu tư thường xuyên đối mặt với tình trạng nhiễu "
        "thông tin, thiếu tính minh bạch và sự chênh lệch về giá cả. Việc ứng dụng các công "
        "nghệ dữ liệu lớn để xây dựng một hệ thống xử lý dòng dữ liệu tập trung, làm sạch và "
        "phân tích xu hướng thị trường không chỉ là một bài toán công nghệ thú vị mà còn mang "
        "lại giá trị thực tiễn cao cho cộng đồng.")

    para(doc,
        "Nhận thức được tầm quan trọng đó, nhóm chúng em đã quyết định thực hiện đề tài: "
        "\"Xây dựng hệ thống tư vấn bất động sản\". Thông qua quá trình thực hiện, nhóm hy vọng "
        "có thể làm rõ quy trình xử lý dữ liệu từ giai đoạn thu thập thô đến khi trực quan hóa, "
        "trả lời yêu cầu tư vấn từ người dùng. Hệ thống gồm các thành phần: thu thập dữ liệu "
        "từ các nguồn trực tuyến, lưu trữ trên cơ sở dữ liệu phân tán và tích hợp tác tử dựa "
        "trên mô hình ngôn ngữ lớn (LLMs agent) để truy xuất thông tin và đưa ra khuyến nghị "
        "đầu tư phù hợp. Quy trình khép kín giúp khai thác dữ liệu hiệu quả, hỗ trợ người dùng "
        "ra quyết định chính xác.")

    para(doc,
        "Đóng góp chính của nhóm là phát triển một hệ thống tích hợp tự động thu thập, xử lý "
        "và lưu trữ dữ liệu đồng thời kết hợp với các mô hình ngôn ngữ lớn để phân tích và hỗ "
        "trợ người dùng tìm kiếm thông tin, từ đó hướng tới việc ra quyết định đầu tư bất "
        "động sản một cách hiệu quả. Phiên bản cập nhật của báo cáo lần này bổ sung thêm các "
        "thành phần Planner Agent và Query Rewriter giúp duy trì ngữ cảnh hội thoại đa lượt, "
        "đồng thời tinh gọn lại giao diện và bộ widget dashboard để tập trung vào các chỉ số "
        "có độ phủ dữ liệu cao trên thực tế.")

    para(doc,
        "Mặc dù đã có nhiều cố gắng trong quá trình nghiên cứu và thực hiện, nhưng do giới hạn "
        "về kiến thức và thời gian, báo cáo chắc chắn không tránh khỏi những thiếu sót. Nhóm "
        "chúng em rất mong nhận được sự đóng góp ý kiến từ thầy và các bạn để đề tài được "
        "hoàn thiện hơn.")

    para(doc, "Chúng em xin chân thành cảm ơn!")
    page_break(doc)


# ============================================================
# MỤC LỤC PLACEHOLDER + DANH MỤC HÌNH/BẢNG
# (Trong Word user phải Update Field để TOC tự sinh — đây chỉ là placeholder)
# ============================================================

def build_toc_placeholder(doc: Document) -> None:
    add_h1(doc, "MỤC LỤC")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(Trong Word: chọn TOC này → nhấn F9 để cập nhật mục lục tự động.)")
    r.font.italic = True

    # Insert TOC field code (Word sẽ render khi user F9)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Mục lục sẽ hiện ở đây sau khi nhấn F9."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)
    page_break(doc)

    add_h1(doc, "DANH MỤC HÌNH VẼ")
    para(doc, "(Cập nhật bằng F9 trong Word.)", indent_first=False)
    bullet(doc, "Hình 3.1: Kiến trúc tổng quan toàn bộ hệ thống")
    bullet(doc, "Hình 3.2: Luồng hoạt động của các tác tử (Planner / SearchDB / Judge / SearchWeb / Writer)")
    bullet(doc, "Hình 3.3: Pipeline Query Rewriter resolving câu hỏi follow-up")
    bullet(doc, "Hình 4.1: Triển khai các instance trên EC2")
    bullet(doc, "Hình 4.2: Luồng giao tiếp của Crawler với Kafka")
    bullet(doc, "Hình 4.3: Airflow lập lịch chạy crawler")
    bullet(doc, "Hình 4.4: Các broker trong cụm Kafka")
    bullet(doc, "Hình 4.5: Log của quá trình xử lý dữ liệu theo luồng (Spark Streaming)")
    bullet(doc, "Hình 4.6: Log của quá trình xử lý dữ liệu theo lô (Spark Batch)")
    bullet(doc, "Hình 4.7: Toàn bộ hệ thống backend đa tác tử (FastAPI + OpenAI + Zep + Langfuse)")
    bullet(doc, "Hình 4.8: Luồng hoạt động của các tác tử trong hệ thống")
    bullet(doc, "Hình 4.9: Danh sách các API trên giao diện FastAPI")
    bullet(doc, "Hình 4.10: Giao diện Chatbot — empty state + multi-chat sidebar")
    bullet(doc, "Hình 4.11: Giao diện Chatbot — turn hội thoại với spinner và follow-up")
    bullet(doc, "Hình 4.12: Dashboard — KPI + Bar giá theo quận")
    bullet(doc, "Hình 4.13: Dashboard — Pie đặc điểm phổ biến theo loại BĐS")
    bullet(doc, "Hình 4.14: Dashboard — Bản đồ nhiệt giá trung bình per tỉnh")
    page_break(doc)

    add_h1(doc, "DANH MỤC BẢNG BIỂU")
    para(doc, "(Cập nhật bằng F9 trong Word.)", indent_first=False)
    bullet(doc, "Bảng 4.1: Coverage % field theo 5 loại BĐS")
    bullet(doc, "Bảng 4.2: Cấu hình widget Dashboard per estate type (DISTRIBUTION_LAYOUT)")
    bullet(doc, "Bảng 5.1: So sánh chi phí và thời gian phản hồi giữa 3 chiến lược Search/Judge")
    bullet(doc, "Bảng 5.2: So sánh hiệu suất phân công của Planner Agent giữa các mô hình")
    bullet(doc, "Bảng 5.3: So sánh hiệu suất gọi công cụ của Search Database Agent")
    bullet(doc, "Bảng 5.4: So sánh hiệu năng giữa các hệ sinh thái LLM (60 requests)")
    bullet(doc, "Bảng 5.5: So sánh hiệu suất phân công giữa các hệ sinh thái")
    bullet(doc, "Bảng 5.6: So sánh hiệu suất gọi công cụ giữa các hệ sinh thái")
    page_break(doc)


# ============================================================
# CHƯƠNG 1 — GIỚI THIỆU ĐỀ TÀI
# ============================================================

def chapter_1(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI")
    doc.add_paragraph()

    add_h2(doc, "1.1  Đặt vấn đề")
    para(doc,
        "Trong thời đại công nghệ số và trí tuệ nhân tạo phát triển mạnh mẽ, dữ liệu đã trở "
        "thành một trong những nguồn tài nguyên quan trọng nhất đối với nhiều lĩnh vực kinh "
        "tế – xã hội. Sự gia tăng nhanh chóng của dữ liệu phi cấu trúc từ các nền tảng trực "
        "tuyến, mạng xã hội và các hệ thống giao dịch điện tử đặt ra yêu cầu cấp thiết về khả "
        "năng thu thập, lưu trữ và phân tích dữ liệu lớn một cách hiệu quả, chính xác và có "
        "khả năng mở rộng.")
    para(doc,
        "Tại Việt Nam, thị trường bất động sản đang phát triển mạnh mẽ trên phạm vi toàn quốc, "
        "trải dài từ các đô thị lớn như Hà Nội, TP. Hồ Chí Minh đến các tỉnh thành đang trong "
        "quá trình đô thị hóa nhanh. Nhu cầu tìm kiếm, giao dịch và đầu tư bất động sản ngày "
        "càng gia tăng, đi kèm với sự bùng nổ của các nền tảng đăng tin trực tuyến. Tuy nhiên, "
        "thị trường này vẫn tồn tại nhiều hạn chế như thông tin phân tán trên nhiều nguồn khác "
        "nhau, thiếu tính minh bạch, dữ liệu không đồng nhất và chất lượng bài đăng không đồng "
        "đều. Những vấn đề này gây khó khăn cho người dùng trong việc tổng hợp thông tin, đánh "
        "giá cơ hội và đưa ra quyết định đầu tư phù hợp.")
    para(doc,
        "Trong bối cảnh đó, việc xây dựng một hệ thống tự động thu thập, lưu trữ và xử lý dữ "
        "liệu bất động sản trên quy mô toàn quốc, dựa trên nền tảng dữ liệu phân tán, kết hợp "
        "với khả năng phân tích và truy xuất thông tin thông minh từ các mô hình ngôn ngữ lớn "
        "(Large Language Models – LLMs), trở nên cần thiết. Hệ thống không chỉ hỗ trợ người "
        "dùng tìm kiếm bất động sản theo nhiều tiêu chí khác nhau mà còn có khả năng phân "
        "tích dữ liệu, tổng hợp thông tin và đưa ra các gợi ý, khuyến nghị nhằm hỗ trợ quá "
        "trình ra quyết định đầu tư.")
    para(doc,
        "Bài toán này đòi hỏi sự tích hợp giữa các lĩnh vực xử lý dữ liệu lớn, trí tuệ nhân "
        "tạo và hệ thống thông minh, mở ra tiềm năng ứng dụng rộng rãi cho thị trường bất "
        "động sản Việt Nam trong bối cảnh chuyển đổi số hiện nay. Xuất phát từ những nhận "
        "định trên, nhóm quyết định lựa chọn đề tài \"Ứng dụng tác tử AI hỗ trợ quyết định "
        "đầu tư bất động sản\" để thực hiện đồ án, với mục tiêu xây dựng một hệ thống có khả "
        "năng mở rộng và áp dụng trên phạm vi toàn quốc.")

    add_h2(doc, "1.2  Các giải pháp hiện tại và hạn chế")
    para(doc,
        "Hiện nay, một số giải pháp đã được áp dụng nhằm hỗ trợ người dùng trong việc tìm "
        "kiếm và ra quyết định đầu tư bất động sản, tuy nhiên vẫn còn tồn tại nhiều hạn chế.")
    bullet(doc,
        "Chatbot trả lời các câu hỏi thường gặp (FAQs): Một số hệ thống sử dụng chatbot với "
        "tập câu hỏi – câu trả lời được xây dựng sẵn để hỗ trợ người dùng. Giải pháp này có "
        "ưu điểm là triển khai đơn giản, phản hồi nhanh nhưng khả năng hiểu ngữ cảnh và xử "
        "lý các câu hỏi phức tạp còn hạn chế. Chatbot khó thích ứng với ngôn ngữ tự nhiên "
        "không chuẩn mực, không có khả năng suy luận sâu hoặc cá nhân hóa tư vấn theo từng "
        "người dùng, dẫn đến hiệu quả hỗ trợ quyết định đầu tư chưa cao.")
    bullet(doc,
        "Các thuật toán học máy và học sâu: Nhiều nghiên cứu và ứng dụng sử dụng các mô hình "
        "học máy, học sâu để phân tích lịch sử để dự đoán hướng giá, nhu cầu thị trường và "
        "các yếu tố ảnh hưởng đến bất động sản. Tuy nhiên, dữ liệu bất động sản tại Việt Nam "
        "thường không đồng nhất, phân tán theo khu vực và biến động nhanh theo thời gian, "
        "đòi hỏi hệ thống phải được cập nhật liên tục. Bên cạnh đó, các mô hình này chủ yếu "
        "tập trung vào dự đoán mà chưa hỗ trợ tốt việc giải thích kết quả hoặc tương tác linh "
        "hoạt với người dùng trong quá trình truy vấn thông tin.")
    para(doc,
        "Nhìn chung, các giải pháp hiện tại vẫn chưa đáp ứng đầy đủ nhu cầu khai thác dữ liệu "
        "bất động sản trên quy mô lớn và hỗ trợ ra quyết định một cách toàn diện. Điều này "
        "đặt ra yêu cầu về một hệ thống thông minh hơn, có khả năng tích hợp dữ liệu đa "
        "nguồn, hiểu ngữ cảnh truy vấn và đưa ra khuyến nghị phù hợp, đóng vai trò như một "
        "tác tử AI hỗ trợ đầu tư bất động sản trên phạm vi toàn quốc.")

    add_h2(doc, "1.3  Mục tiêu và phạm vi của dự án")
    para(doc,
        "Về mục tiêu, dự án nhằm xây dựng một hệ thống có khả năng tự động thu thập, lưu trữ, "
        "xử lý và phân tích hiệu quả dữ liệu liên quan đến bất động sản trên cả nước. Mục "
        "tiêu và định hướng cụ thể bao gồm:")
    numbered(doc,
        "Xây dựng thành phần lập lịch tự động thu thập dữ liệu thời gian thực. Dữ liệu sẽ "
        "được xử lý và lưu trữ trên các nền tảng phân tán phục vụ khả năng mở rộng.")
    numbered(doc,
        "Thiết kế kiến trúc đa tác tử dựa trên dữ liệu đã thu thập được. Các tác tử này sẽ "
        "sử dụng mô hình ngôn ngữ lớn (LLMs) làm bộ não suy luận để tổng hợp thông tin về "
        "bất động sản và đưa phân tích dựa trên câu hỏi của người dùng. Phiên bản cập nhật "
        "bổ sung thành phần Query Rewriter cho phép duy trì ngữ cảnh hội thoại đa lượt.")
    numbered(doc,
        "Xây dựng giao diện web gồm: giao diện hỏi đáp chatbot và các biểu đồ trực quan hóa "
        "dữ liệu thiết kế theo nhóm người dùng (người mua / nhà đầu tư).")
    para(doc, "Về phạm vi, dự án tập trung vào xây dựng hệ thống trong đó:")
    bullet(doc, "Đối tượng: Các bài đăng bán bất động sản nhà riêng, chung cư, biệt thự, nhà mặt phố, đất đai và phân loại \"khác\".")
    bullet(doc, "Khu vực: Các quận, huyện thuộc 63 tỉnh thành tại Việt Nam.")
    bullet(doc, "Thời gian dữ liệu: Thu thập và phân tích dữ liệu trong giai đoạn cuối năm 2025 đến đầu năm 2026, với khoảng 87.000 bản ghi thực tế tại thời điểm viết báo cáo cập nhật.")

    add_h2(doc, "1.4  Bố cục của dự án")
    para(doc, "Phần còn lại của dự án này được tổ chức như sau.")
    para(doc,
        "Chương 2 - Nền tảng lý thuyết trình bày về ngữ cảnh của bài toán, các kết quả nghiên "
        "cứu tương tự liên quan đến việc ứng dụng chatbot và agent dựa trên mô hình ngôn ngữ "
        "lớn (LLMs) trong ngành bất động sản. Nội dung bao gồm kiến trúc và nguyên lý của các "
        "nền tảng xử lý dữ liệu phân tán như Kafka và Spark, giới thiệu về mô hình ngôn ngữ "
        "lớn (LLMs) và tác tử LLMs cùng như các kỹ thuật prompt và luồng hoạt động của các "
        "tác tử.")
    para(doc, "Chương 3 - Phương pháp đề xuất: mô tả chi tiết các giải pháp đề xuất dựa trên các mục tiêu và định hướng của dự án bao gồm:")
    numbered(doc, "Thực hiện tự động lập lịch thu thập và xử lý dữ liệu, lưu trữ dữ liệu trên hệ thống phân tán.")
    numbered(doc, "Kết hợp dữ liệu được lưu trữ để thiết kế một kiến trúc đa tác tử để truy xuất thông tin, gồm Planner, Search DB, Search Web, Judge, Writer và Query Rewriter.")
    numbered(doc, "Xây dựng một ứng dụng web tương tác chatbot và các biểu đồ trực quan hóa dữ liệu được thiết kế theo từng phân khúc người dùng.")
    para(doc,
        "Chương 4 - Triển khai hệ thống: Trình bày về quá trình xây dựng và triển khai thực "
        "tế các giải pháp đề xuất để tạo nên hệ thống, từ việc thu thập dữ liệu, xử lý dữ "
        "liệu đến việc thiết kế hệ thống đa tác tử dựa trên OpenAI Python SDK và phát triển "
        "ứng dụng web Streamlit cho phép người dùng tương tác với hệ thống.")
    para(doc,
        "Chương 5 - Đánh giá thực nghiệm: trình bày các kết quả thí nghiệm thu được khi áp "
        "dụng hệ thống trên các tập dữ liệu khác nhau. Nội dung bao gồm: Mô tả các tham số "
        "thí nghiệm, phương pháp tiến hành thí nghiệm và các kết quả kèm theo nhận xét, đặc "
        "biệt là so sánh giữa hai hệ sinh thái Gemini và OpenAI sau khi chuyển stack chính "
        "thức sang OpenAI.")
    para(doc,
        "Chương 6 - Kết luận và hướng phát triển: tóm tắt các kết quả đạt được của dự án, "
        "nhấn mạnh những đóng góp chính và những hạn chế còn tồn tại. Phần này cũng đề xuất "
        "những hướng phát triển trong tương lai để cải thiện và mở rộng nghiên cứu.")
    page_break(doc)


# ============================================================
# CHƯƠNG 2 — NỀN TẢNG LÝ THUYẾT
# (Giữ gần như nguyên report cũ vì là lý thuyết tổng quan)
# ============================================================

def chapter_2(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 2. NỀN TẢNG LÝ THUYẾT")
    doc.add_paragraph()

    add_h2(doc, "2.1  Ngữ cảnh của bài toán")
    para(doc,
        "Thị trường bất động sản tại Việt Nam đang ngày càng phát triển với nhiều biến động "
        "phức tạp, kéo theo nhu cầu cấp thiết trong việc thu thập, lưu trữ và phân tích dữ "
        "liệu từ nhiều nguồn khác nhau như sàn giao dịch, mạng xã hội, các trang tin tức và "
        "website rao bán nhà đất. Việc xử lý hiệu quả khối lượng lớn dữ liệu này theo thời "
        "gian thực không chỉ giúp cập nhật nhanh chóng tình hình thị trường mà còn đóng vai "
        "trò quan trọng trong việc hỗ trợ người dùng tìm kiếm, đánh giá và đưa ra quyết định "
        "đầu tư hoặc mua bán một cách chính xác và kịp thời.")
    para(doc,
        "Bên cạnh đó, việc tích hợp các mô hình ngôn ngữ lớn (LLMs) dưới dạng các tác tử "
        "thông minh (LLM Agents) vào hệ thống còn mở ra khả năng hiểu và phản hồi linh hoạt "
        "các truy vấn của người dùng liên quan đến bất động sản. Các tác vụ như tìm kiếm "
        "căn nhà phù hợp theo tiêu chí cá nhân, so sánh giá bất động sản giữa các khu vực, "
        "phân tích tiềm năng đầu tư của một địa điểm, hay dự đoán xu hướng thị trường đều có "
        "thể được hỗ trợ hiệu quả nhờ sự phối hợp giữa hệ thống dữ liệu thời gian thực và "
        "các LLM Agents.")

    add_h2(doc, "2.2  Các kết quả nghiên cứu tương tự")
    para(doc, "Nhiều nghiên cứu đã được thực hiện trong lĩnh vực ứng dụng chatbot và agent dựa trên mô hình ngôn ngữ lớn (LLMs) trong ngành bất động sản. Một số nghiên cứu tiêu biểu bao gồm:")
    bullet(doc,
        "Nghiên cứu sử dụng chatbot dựa trên quy tắc (rule-based chatbots): Các chatbot này "
        "sử dụng các kịch bản được định nghĩa trước để trả lời các câu hỏi thường gặp về bất "
        "động sản, như thông tin giá cả hoặc lịch xem nhà. Ưu điểm của phương pháp này là dễ "
        "triển khai và chi phí thấp, phù hợp cho các doanh nghiệp nhỏ. Tuy nhiên, nhược điểm "
        "là khả năng xử lý các truy vấn phức tạp hoặc không nằm trong kịch bản bị hạn chế, "
        "dẫn đến trải nghiệm người dùng kém khi gặp các câu hỏi mang tính ngữ cảnh hoặc cá "
        "nhân hóa.")
    bullet(doc,
        "Nghiên cứu sử dụng học máy và NLP cơ bản: Các nghiên cứu này áp dụng các thuật toán "
        "học máy kết hợp với xử lý ngôn ngữ tự nhiên (NLP) để xây dựng chatbot có khả năng "
        "phân tích và trả lời các câu hỏi của khách hàng về bất động sản, chẳng hạn như tìm "
        "kiếm nhà theo ngân sách hoặc vị trí. Phương pháp này cải thiện khả năng hiểu ngữ "
        "cảnh so với chatbot dựa trên quy tắc, nhưng vẫn gặp khó khăn trong việc xử lý các "
        "truy vấn phức tạp hoặc dữ liệu thời gian thực, ví dụ như cập nhật giá nhà mới nhất "
        "hoặc thông tin thị trường. Ngoài ra, các hệ thống này yêu cầu dữ liệu huấn luyện "
        "lớn và thường thiếu tính cá nhân hóa.")

    add_h2(doc, "2.3  Kiến trúc và nguyên lý hoạt động của Apache Kafka")
    para(doc,
        "Apache Kafka là một nền tảng xử lý dòng sự kiện phân tán (Distributed Event "
        "Streaming Platform) được phát triển bởi Apache Software Foundation. Kafka được "
        "thiết kế để xử lý, lưu trữ và truyền tải một lượng lớn thông điệp theo thời gian "
        "thực từ nhiều nguồn dữ liệu khác nhau, với độ trễ thấp và khả năng mở rộng cao.")

    add_h3(doc, "2.3.1  Mô hình lưu trữ Log-structured")
    para(doc,
        "Kafka sử dụng mô hình lưu trữ dạng log, trong đó dữ liệu được tổ chức thành các "
        "dòng thông điệp (Messages) được ghi tuần tự vào đĩa. Mỗi thông điệp có thể chứa "
        "một hoặc nhiều byte dữ liệu, và được ghi vào các phân vùng (Partitions) theo thứ "
        "tự thời gian. Dữ liệu chỉ được ghi nối tiếp (append-only), không hỗ trợ chỉnh sửa "
        "hay xóa.")

    add_h3(doc, "2.3.2  Mô hình Pub-Sub với Topic và Partition")
    para(doc, "Kafka vận hành theo mô hình Publisher - Subscriber, trong đó:")
    bullet(doc, "Producer (Nhà cung cấp dữ liệu): gửi thông điệp vào một hoặc nhiều chủ đề (Topic).")
    bullet(doc, "Consumer (Nhà tiêu thụ dữ liệu): đăng ký nhận thông điệp từ các Topic tương ứng.")
    para(doc,
        "Mỗi Topic được chia thành nhiều Partition để hỗ trợ song song hóa và phân tán dữ "
        "liệu. Thông điệp trong mỗi Partition được sắp xếp theo thứ tự và đánh dấu bằng một "
        "chỉ số duy nhất gọi là Offset.")

    add_h3(doc, "2.3.3  Cấu trúc Cluster và các thành phần chính")
    para(doc, "Kafka bao gồm nhiều thành phần hoạt động phối hợp trong một cụm (Kafka Cluster), với các thành phần chính như sau:")
    bullet(doc, "Broker: Máy chủ lưu trữ và xử lý các Partition. Mỗi broker có thể chứa nhiều Partition của các Topic khác nhau.")
    bullet(doc, "ZooKeeper: Quản lý metadata, điều phối các broker, theo dõi tình trạng leader và phân vùng.")
    bullet(doc, "Topic: Chủ đề chứa các luồng dữ liệu được phân loại.")
    bullet(doc, "Producer / Consumer: Gửi và nhận dữ liệu từ Kafka.")
    para(doc, "Kiến trúc phân tán của Kafka cho phép mở rộng bằng cách thêm broker vào cluster mà không làm gián đoạn hoạt động của hệ thống.")

    add_h3(doc, "2.3.4  Cơ chế phân phối và sao lưu dữ liệu")
    para(doc,
        "Để đảm bảo độ tin cậy và khả năng chịu lỗi, Kafka hỗ trợ cơ chế Replication — mỗi "
        "Partition có thể được sao chép sang nhiều broker khác nhau. Trong số các bản sao, "
        "một broker giữ vai trò Leader, các broker còn lại là Follower. Tất cả các thao tác "
        "đọc/ghi đều được thực hiện qua Leader, và các Follower đồng bộ hóa dữ liệu để đảm "
        "bảo tính sẵn sàng trong trường hợp sự cố xảy ra.")

    add_h2(doc, "2.4  Xử lý dữ liệu lớn với Apache Spark")
    para(doc,
        "Apache Spark là một nền tảng xử lý dữ liệu phân tán được sử dụng để xử lý và phân "
        "tích dữ liệu lớn. Nó được thiết kế để xử lý dữ liệu lớn và phức tạp bằng cơ chế "
        "in-memory và lazy evaluation, xử lý dữ liệu trên bộ nhớ trong và chỉ thực hiện các "
        "thao tác độc ghi khi cần, giúp cung cấp tốc độ xử lý nhanh chóng và khả năng mở "
        "rộng linh hoạt. Spark hỗ trợ nhiều ngôn ngữ lập trình như Scala, Java, Python và R, "
        "thuận tiện cho việc tích hợp và triển khai các ứng dụng xử lý dữ liệu phức tạp.")

    add_h3(doc, "2.4.1  Giới thiệu về Spark Streaming")
    para(doc,
        "Spark Streaming là một thành phần mở rộng trong hệ sinh thái Apache Spark, được "
        "thiết kế để xử lý dữ liệu dòng (streaming data) — tức dữ liệu liên tục được tạo ra "
        "theo thời gian thực từ các nguồn như cảm biến, mạng xã hội, hệ thống giao dịch, "
        "hoặc log hệ thống. Không giống như mô hình xử lý theo lô truyền thống, Spark "
        "Streaming cho phép xử lý dữ liệu gần thời gian thực bằng một kiến trúc được tối ưu "
        "hóa cho khả năng mở rộng và hiệu suất cao.")

    add_h3(doc, "2.4.2  Mô hình xử lý Micro-Batch")
    para(doc,
        "Spark Streaming áp dụng mô hình micro-batch, trong đó dòng dữ liệu được chia thành "
        "các lô nhỏ (mini-batches) theo chu kỳ thời gian cố định (ví dụ: mỗi 1 giây hoặc 5 "
        "giây). Mỗi lô dữ liệu nhỏ này sau đó được xử lý như một batch thông thường bằng "
        "Spark Engine.")

    add_h3(doc, "2.4.3  Kiến trúc dòng xử lý DStream")
    para(doc,
        "Spark Streaming xây dựng dòng xử lý dựa trên cấu trúc DStream (Discretized Stream). "
        "Một DStream đại diện cho một chuỗi các RDD (Resilient Distributed Dataset) được "
        "tạo ra theo thời gian. Mỗi RDD trong DStream tương ứng với một lô dữ liệu thu thập "
        "được trong khoảng thời gian cụ thể.")

    add_h3(doc, "2.4.4  Khả năng mở rộng và phục hồi lỗi")
    para(doc,
        "Spark Streaming hỗ trợ khả năng mở rộng ngang (scale-out) nhờ vào kiến trúc phân "
        "tán của Spark. Việc mở rộng hệ thống đơn giản chỉ bằng cách thêm node vào cluster "
        "Spark. Đồng thời, nhờ cơ chế RDD, hệ thống có khả năng khôi phục khi gặp lỗi "
        "(fault-tolerant) thông qua việc ghi log các thao tác chuyển đổi (lineage) để có "
        "thể tính toán lại nếu xảy ra sự cố.")

    add_h2(doc, "2.5  Tác tử AI dựa trên mô hình ngôn ngữ lớn")

    add_h3(doc, "2.5.1  Giới thiệu về Mô hình ngôn ngữ lớn (LLMs)")
    para(doc,
        "Large Language Models (LLMs) là các mô hình học sâu được huấn luyện trên khối "
        "lượng văn bản khổng lồ nhằm học cách biểu diễn và sinh ngôn ngữ tự nhiên. Các mô "
        "hình lớn như GPT (Generative Pretrained Transformer) của OpenAI, Gemini của Google "
        "hay Llama của Meta đã đạt được những bước tiến vượt bậc trong việc hiểu và sinh "
        "ngôn ngữ, đồng thời trở thành nền tảng cho nhiều ứng dụng trong xử lý ngôn ngữ tự "
        "nhiên (NLP).")
    para(doc,
        "LLMs hoạt động dựa trên kiến trúc Transformer, cho phép mô hình nắm bắt ngữ cảnh "
        "dài và tạo ra các phản hồi phù hợp theo yêu cầu người dùng. Việc huấn luyện mô "
        "hình thường sử dụng hai giai đoạn: huấn luyện không có giám sát (unsupervised "
        "pretraining) và tinh chỉnh (fine-tuning) với dữ liệu có giám sát hoặc thông qua "
        "Reinforcement Learning from Human Feedback (RLHF).")

    add_h3(doc, "2.5.2  Prompt Engineering và các kỹ thuật prompt")
    para(doc,
        "Trong xử lý ngôn ngữ tự nhiên, Prompt là một câu hỏi hoặc một đoạn văn bản dẫn dắt "
        "mà con người sử dụng để tương tác với các mô hình AI. Prompt Engineering là quá "
        "trình tạo ra đầu vào cho các mô hình AI, thiết kế Prompt một cách tỉ mỉ và logic "
        "để có thể kiểm soát được kết quả đầu ra từ các mô hình AI một cách chính xác. Dưới "
        "đây là một số kỹ thuật prompt phổ biến:")
    bold_label(doc, "Zero-shot prompting:",
        "Yêu cầu mô hình thực hiện một tác vụ mà không cần cung cấp bất kỳ ví dụ nào trước đó. "
        "Nó dựa vào khả năng suy luận và kiến thức đã được học trong quá trình huấn luyện mô hình.")
    bold_label(doc, "Few-shot prompting:",
        "Cung cấp một vài ví dụ mẫu để mô hình học cách thực hiện tác vụ. Các ví dụ này giúp "
        "mô hình suy luận theo đúng định dạng hoặc quy tắc mong muốn.")
    bold_label(doc, "Chain-of-thought prompting (CoT):",
        "Kỹ thuật này khuyến khích mô hình giải thích suy luận từng bước để đến kết quả, đặc "
        "biệt hữu ích trong các tác vụ toán học, logic, hoặc ra quyết định.")
    bold_label(doc, "Role prompting:",
        "Mô hình được giao một vai trò cụ thể (ví dụ: giáo viên, lập trình viên, bác sĩ) để "
        "định hình phong cách phản hồi phù hợp với ngữ cảnh. Trong dự án này, Writer Agent "
        "được role-prompt thành \"nhà phân tích bất động sản\" để output có cấu trúc tư vấn.")

    add_h3(doc, "2.5.3  Các ứng dụng chính của LLMs")
    bullet(doc, "Tóm tắt văn bản, dịch máy, sinh mã nguồn (Code Generation).")
    bullet(doc, "Hỏi đáp và truy xuất thông tin: LLMs có khả năng tìm kiếm, phân tích và tổng hợp thông tin từ các kho dữ liệu lớn, cung cấp câu trả lời chính xác và chi tiết cho các câu hỏi phức tạp — ứng dụng chính trong đề tài này.")

    add_h3(doc, "2.5.4  Kiến trúc của tác tử LLMs")
    para(doc,
        "Tác tử LLMs là sự kết hợp giữa các mô hình ngôn ngữ lớn với các thành phần hỗ trợ "
        "nhằm tạo ra các tác tử (agent) có khả năng hành động, suy luận và tương tác tự "
        "động với môi trường. Không giống như LLMs truyền thống (chỉ phản hồi đầu vào), tác "
        "tử LLMs có khả năng:")
    bullet(doc, "Lập kế hoạch (Planning): xây dựng các kế hoạch để giải quyết các nhiệm vụ phức tạp.")
    bullet(doc, "Gọi công cụ (Tool Usage): tương tác với các API, cơ sở dữ liệu hoặc phần mềm chuyên dụng để truy cập kiến thức bên ngoài.")
    bullet(doc, "Ghi nhớ trạng thái (Memory): lưu trữ và truy xuất thông tin từ các tương tác trước, duy trì ngữ cảnh liên tục qua nhiều phiên làm việc.")
    bullet(doc, "Học và phản hồi từ kết quả (Feedback Loop): học hỏi từ phản hồi để liên tục cải thiện hiệu suất qua mỗi vòng lặp.")
    bold_label(doc, "Thành phần chính của tác tử LLMs:", "")
    bullet(doc, "LLMs Engine: Bộ não của agent, thực hiện sinh văn bản, suy luận logic và kết nối giữa các bước.")
    bullet(doc, "Memory Module: short-term (prompt window) + long-term (Vector DB, Knowledge Graph như ZepAI).")
    bullet(doc, "Tool Use (External Calls): cho phép agent gọi các API ngoài, tìm kiếm thông tin trên web, truy vấn cơ sở dữ liệu.")
    bold_label(doc, "Luồng hoạt động phổ biến của hệ thống đa tác tử:", "")
    bullet(doc, "Prompt chaining: chia nhiệm vụ thành chuỗi bước, mỗi lần gọi LLM xử lý đầu ra của bước trước đó.")
    bullet(doc, "Routing: phân loại đầu vào và chuyển tới các LLM/agent chuyên biệt — đây chính là vai trò của Planner Agent trong dự án này.")
    bullet(doc, "Evaluator-Optimizer: dùng hai mô hình LLM phối hợp, một sinh đầu ra (generator), một đánh giá (evaluator) để cải thiện chất lượng. Đây là pattern Judge Agent áp dụng trong vòng self-correction với Search Database Agent.")

    add_h2(doc, "2.6  Kết luận")
    para(doc,
        "Trong chương này, chúng ta đã có cái nhìn tổng quan về bối cảnh và ngữ cảnh của bài "
        "toán ứng dụng chatbot và agent dựa trên mô hình ngôn ngữ lớn (LLMs) trong lĩnh vực "
        "bất động sản, đồng thời khảo sát các nghiên cứu liên quan. Chúng ta cũng đã điểm "
        "qua những lý thuyết nền tảng để xây dựng một hệ thống toàn diện từ thu thập và lưu "
        "trữ dữ liệu với Kafka và Spark đến tìm hiểu kiến trúc tác tử LLMs để hoàn thành các "
        "tác vụ phức tạp mà người dùng yêu cầu. Trong phần tiếp theo, dự án sẽ đề xuất các "
        "phương pháp để hiện thực hóa nền tảng lý thuyết trên.")
    page_break(doc)


# ============================================================
# Phần Chương 3-6 sẽ được module khác hoặc tiếp tục mở rộng
# ============================================================

def chapter_3(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 3. PHƯƠNG PHÁP ĐỀ XUẤT")
    doc.add_paragraph()

    add_h2(doc, "3.1  Tổng quan giải pháp")
    para(doc,
        "Trong bài tập lớn này, nhóm chúng em đề xuất một giải pháp tích hợp từ việc xây "
        "dựng pipeline thu thập – lưu trữ – xử lý dữ liệu đến thiết kế kiến trúc đa tác tử "
        "để thực hiện các tác vụ mà người dùng yêu cầu và một giao diện web thân thiện, hỗ "
        "trợ nhà đầu tư và người mua bất động sản tương tác với dữ liệu thông qua chatbot "
        "và các biểu đồ trực quan. Hình 3.1 mô tả quy trình tổng quan của hệ thống, gồm các "
        "thành phần:")
    bold_label(doc, "Quy trình thu thập, xử lý và lưu trữ dữ liệu:",
        "Dữ liệu được lập lịch để tự động crawl từ các bài đăng bán, cho thuê từ các website "
        "bất động sản. Dữ liệu sau đó được truyền tải qua các nền tảng phân tán để xử lý và "
        "chuẩn hóa, cuối cùng được lưu vào cơ sở dữ liệu nhằm phục vụ truy vấn.")
    bold_label(doc, "Kiến trúc đa tác tử:",
        "Hệ thống xây dựng một kiến trúc đa tác tử nhằm tận dụng nguồn dữ liệu thu thập được "
        "lưu trữ trong cơ sở dữ liệu để tìm kiếm các bất động sản theo yêu cầu của người "
        "dùng, đồng thời kết nối với dữ liệu trên Internet để tìm các thông tin liên quan "
        "đến pháp lý, tiện ích xung quanh hay tin tức thị trường. Phiên bản cập nhật bổ "
        "sung Query Rewriter giúp tác tử duy trì ngữ cảnh đa lượt.")
    bold_label(doc, "Ứng dụng web:",
        "Xây dựng giao diện tương tác với người dùng gồm 2 chức năng chính: hỏi đáp chatbot "
        "và các biểu đồ trực quan hóa dữ liệu. Việc xử lý yêu cầu của người dùng được kết "
        "nối với kiến trúc các tác tử và cơ sở dữ liệu nhằm lấy được những thông tin để "
        "biểu diễn trên giao diện.")

    add_h2(doc, "3.2  Đề xuất giải pháp thu thập, tiền xử lý và lưu trữ dữ liệu")
    add_h3(doc, "3.2.1  Thu thập và tiền xử lý dữ liệu")
    para(doc,
        "Trong bài toán xây dựng hệ thống tư vấn bất động sản dựa trên dữ liệu lớn, thành "
        "phần thu thập và tiền xử lý dữ liệu đóng vai trò nền tảng, quyết định trực tiếp đến "
        "chất lượng dữ liệu đầu vào cũng như hiệu quả của các mô hình phân tích và tư vấn ở "
        "các giai đoạn tiếp theo. Do dữ liệu bất động sản trên các nền tảng trực tuyến "
        "thường có tính phân tán, không đồng nhất và chứa nhiều nhiễu, việc thiết kế một "
        "giải pháp thu thập và tiền xử lý tự động, linh hoạt và có khả năng mở rộng là yêu "
        "cầu bắt buộc.")
    para(doc,
        "Giải pháp được thiết kế gồm ba thành phần chính: (1) dịch vụ thu thập dữ liệu, "
        "(2) pipeline tiền xử lý dữ liệu và (3) dịch vụ lập lịch và điều phối. Các thành "
        "phần này được triển khai độc lập nhưng liên kết chặt chẽ với nhau thông qua các "
        "giao diện cấu hình và cơ chế truyền dữ liệu chuẩn hóa.")
    para(doc,
        "Dữ liệu được thu thập từ các bài đăng bán bất động sản trên các website phổ biến "
        "tại Việt Nam (cụ thể là batdongsan68.com.vn và bds68.com.vn), sử dụng framework mã "
        "nguồn mở Scrapy để triển khai các Spider thu thập dữ liệu. Scrapy gồm các thành phần "
        "Spider (định nghĩa cách truy cập URL, trích xuất thông tin) và Items (schema các "
        "trường thông tin cần trích rút).")
    bold_label(doc, "Các trường được trích xuất:", "")
    for f in ["title — tiêu đề bài đăng", "description — nội dung mô tả",
              "price — giá bất động sản", "square — diện tích",
              "estate_type — loại hình BĐS (nhà mặt phố, nhà riêng, chung cư, biệt thự, đất, khác)",
              "address — địa chỉ bài đăng (province, district, ward)",
              "post_date — ngày đăng bài",
              "contact_info — thông tin liên hệ (name, phone[])",
              "extra_infos — số tầng / số phòng ngủ / số phòng tắm / mặt tiền / năm xây dựng / hướng nhà",
              "link — liên kết gốc của bài đăng"]:
        bullet(doc, f)
    bold_label(doc, "Pipeline tiền xử lý dữ liệu:",
        "Sau khi trích xuất, dữ liệu thô được đưa qua pipeline tiền xử lý gồm ba bước: hiệu "
        "chỉnh chuẩn hóa địa chỉ (loại bỏ tiền tố \"Tỉnh/Thành phố/Quận/Phường\", đối chiếu "
        "với tập đơn vị hành chính trong JSON, fuzzy matching qua thư viện geoapivietnam); "
        "loại bỏ bài đăng trùng lặp (TF-IDF + cosine similarity với ngưỡng 0.95/0.99); và "
        "đẩy vào Kafka topic tương ứng.")

    add_h3(doc, "3.2.2  Lưu trữ và xử lý dữ liệu")
    para(doc,
        "Đặc thù dữ liệu bất động sản là tăng trưởng liên tục theo thời gian, dữ liệu bài "
        "đăng có tính lịch sử và nhiều trường văn bản dài (title, description), đồng thời "
        "cần hỗ trợ truy vấn theo từ khóa và theo thuộc tính (giá, diện tích, vị trí). Vì "
        "vậy, hệ thống sử dụng mô hình hybrid storage, kết hợp giữa lưu trữ lâu dài và lưu "
        "trữ tối ưu truy vấn:")
    bullet(doc, "Ingestion Layer (Kafka): tiếp nhận dữ liệu từ Crawler, đóng vai trò hàng đợi thông điệp tin cậy.")
    bullet(doc, "Storage Layer (MinIO): lưu trữ dài hạn dữ liệu thô và đã xử lý dưới dạng data lake (Parquet).")
    bullet(doc, "Serving Layer (Elasticsearch): cung cấp truy vấn nhanh, tìm kiếm văn bản và aggregation cho dashboard.")
    para(doc,
        "Hệ thống kết hợp hai mô hình xử lý: stream processing (Spark Structured Streaming "
        "đọc từ Kafka topic và ghi xuống Elasticsearch + MinIO) cho dữ liệu thời gian thực, "
        "và batch processing (Spark Batch Job định kỳ qua Airflow) cho phân tích dữ liệu "
        "lịch sử, recompute aggregate, loại bỏ outlier và cập nhật ngược ES.")

    add_h2(doc, "3.3  Đề xuất giải pháp Kiến trúc Đa tác tử dựa trên Mô hình Ngôn ngữ Lớn")
    para(doc,
        "Hệ thống tư vấn bất động sản dựa trên mô hình ngôn ngữ lớn (LLMs) sử dụng kiến "
        "trúc đa tác tử (multi-agent architecture), trong đó mỗi tác tử (agent) đảm nhận một "
        "nhiệm vụ riêng biệt và có thể tương tác với nhau để đạt được mục tiêu chung. Mô "
        "hình này không chỉ tăng cường khả năng mở rộng mà còn tạo ra một hệ thống linh "
        "hoạt, có khả năng phân chia công việc một cách hiệu quả. Cụ thể, kiến trúc đa tác "
        "tử được xây dựng trên ba mẫu chính: Routing (Planner Agent quyết định tool), Prompt "
        "Chaining (Rewriter → Planner → Search → Writer), và Evaluator-Optimizer (Judge "
        "self-correction loop với Search Database Agent).")
    para(doc,
        "Mô hình hoạt động của hệ thống và các tác tử được mô tả trong Hình 3.2. Hệ thống "
        "bao gồm sáu tác tử chính (so với bản cũ 5 tác tử — phiên bản này bổ sung Query "
        "Rewriter), mỗi tác tử được thiết kế để thực hiện một nhiệm vụ cụ thể và có thể phối "
        "hợp với nhau để cung cấp thông tin chính xác và kịp thời cho người dùng:")
    bullet(doc, "Query Rewriter — viết lại câu hỏi theo ngữ cảnh hội thoại (MỚI).")
    bullet(doc, "Planner Agent — phân loại intent và lựa chọn tool (tên cũ \"Manager Agent\").")
    bullet(doc, "Search Database Agent — truy vấn Elasticsearch tìm bài đăng phù hợp.")
    bullet(doc, "Search Web Agent — tìm thông tin trên Internet (DuckDuckGo).")
    bullet(doc, "Judge Agent — đánh giá kết quả search, vòng lặp self-correction.")
    bullet(doc, "Writer Agent — tổng hợp dữ liệu thành báo cáo tư vấn có cấu trúc.")

    add_h3(doc, "3.3.1  Query Rewriter Agent (Mới)")
    para(doc,
        "Trong các cuộc hội thoại đa lượt, người dùng thường gửi các câu hỏi follow-up phụ "
        "thuộc ngữ cảnh, ví dụ \"với 50 tỷ thì mua được căn nào\" sau khi đã hỏi về \"nhà mặt "
        "phố Hà Nội Hoàn Kiếm\". Nếu chỉ truyền câu hỏi mới cho Search Agent, các bộ lọc đã "
        "nói ở turn trước sẽ bị mất, dẫn đến kết quả truy vấn sai phạm vi.")
    para(doc, "Query Rewriter Agent giải quyết vấn đề này bằng cách:")
    bullet(doc, "Đọc 16 message gần nhất (8 turn user + assistant) làm bối cảnh.")
    bullet(doc, "Áp dụng prompt few-shot có 4 ví dụ resolve các tình huống follow-up phổ biến (bổ sung filter cũ, refer item theo số thứ tự, refer theo tính chất).")
    bullet(doc, "Trả về câu hỏi standalone đầy đủ ngữ cảnh, vd: \"Mua nhà mặt phố Hà Nội quận Hoàn Kiếm với ngân sách 50 tỷ thì mua được căn nào\".")
    para(doc,
        "Câu standalone sau khi rewrite được dùng làm đầu vào cho tất cả các tác tử phía "
        "sau (Planner, SearchDB, Writer), đảm bảo mọi tác tử đều thấy ngữ cảnh đầy đủ. "
        "Nếu câu hỏi đã standalone hoặc không phụ thuộc lịch sử, Rewriter trả về nguyên văn.")

    add_h3(doc, "3.3.2  Planner Agent (đổi tên từ Manager Agent)")
    para(doc,
        "Planner Agent là tác tử đầu tiên xử lý câu hỏi sau khi đã được rewrite. Nhiệm vụ "
        "của agent này gồm hai phần: (A) phân loại intent của câu hỏi vào một trong ba "
        "nhóm — bds_query (câu hỏi liên quan bất động sản), chitchat (xã giao, hỏi bot là "
        "ai), off_topic (câu lạc đề); và (B) lựa chọn tool sẽ được kích hoạt khi "
        "intent=bds_query (search_db, search_web, hoặc cả hai). Agent được hướng dẫn bằng "
        "prompt dạng few-shot với 9 ví dụ phủ các tình huống biên.")
    para(doc,
        "Với intent thuộc chitchat hoặc off_topic, Planner Agent kích hoạt cơ chế "
        "short-circuit — không gọi Search/Writer, mà trả về response cố định "
        "(\"Xin chào, mình là trợ lý BĐS...\" hoặc lời từ chối lịch sự kèm gợi ý câu hỏi mẫu). "
        "Cơ chế này tiết kiệm chi phí token và cải thiện đáng kể độ trễ cho các câu hỏi không "
        "phải BĐS.")

    add_h3(doc, "3.3.3  Search Database Agent")
    para(doc,
        "Search Database Agent chịu trách nhiệm tìm kiếm các bài đăng bất động sản phù hợp "
        "với yêu cầu của người dùng từ Elasticsearch. Tác tử này có hai hàm tìm kiếm chính:")
    bold_label(doc, "search_posts_strict:",
        "Tìm kiếm các bài đăng thỏa mãn đồng thời tất cả các tiêu chí mà người dùng đưa ra "
        "(địa chỉ, mức giá ±10%, diện tích, số phòng ngủ, số phòng tắm…). Phù hợp với query "
        "cụ thể có nhiều ràng buộc rõ ràng.")
    bold_label(doc, "search_posts (loose):",
        "Tổng hợp các bài đăng thỏa mãn từng tiêu chí riêng lẻ, với phạm vi nới rộng ±20% "
        "cho giá/diện tích và text-fallback cho địa điểm. Trả về tối đa 12 bài không trùng "
        "lặp. Phù hợp với query mơ hồ.")
    para(doc,
        "Chiến lược của tác tử là ưu tiên dùng search_posts_strict khi câu hỏi có ≥2 ràng "
        "buộc rõ ràng (vd: tỉnh + giá + loại). Nếu strict trả về rỗng, tự động fallback "
        "sang search_posts (loose). Việc lựa chọn được điều phối qua prompt few-shot, giúp "
        "tác tử tự động suy luận tham số khi gọi các hàm tìm kiếm.")

    add_h3(doc, "3.3.4  Search Web Agent")
    para(doc,
        "Search Web Agent có nhiệm vụ tìm kiếm các thông tin trên Internet liên quan đến "
        "các thông tin mà cơ sở dữ liệu không có như xu hướng và dự báo thị trường, pháp lý "
        "và các tiện ích xung quanh. Phiên bản cập nhật chuyển từ Serper API (yêu cầu API "
        "key trả phí) sang thư viện DuckDuckGo Search (ddgs) — miễn phí và không cần API "
        "key. Agent gọi DuckDuckGo lấy top 5 kết quả, sau đó dùng OpenAI để tóm tắt thành "
        "summary 2-3 đoạn dưới 300 từ.")

    add_h3(doc, "3.3.5  Judge Agent")
    para(doc, "Judge Agent đánh giá lại chất lượng kết quả từ Search Database Agent. Quy trình đánh giá:")
    numbered(doc, "Xác định mức độ liên quan của từng kết quả đến câu hỏi người dùng (cho phép ±10% cho price/square/price_per_square, các trường khác phải khớp chính xác trừ description fuzzy).")
    numbered(doc, "Nếu ≥30% kết quả đạt tiêu chí — score=pass, gửi cho Writer Agent.")
    numbered(doc, "Nếu <30% — score=needs_improvement, kèm feedback hướng dẫn Search Agent rewrite query và lặp lại (tối đa 2 vòng).")
    para(doc, "Output Judge gồm 3 trường: feedback, score (pass/needs_improvement), reason. Khi Judge gặp lỗi LLM, fallback trả \"pass\" để pipeline không tắc.")

    add_h3(doc, "3.3.6  Writer Agent")
    para(doc,
        "Writer Agent là tác tử cuối cùng trong chuỗi xử lý, được hướng dẫn bằng Role "
        "prompt trong vai trò một nhà phân tích bất động sản chuyên nghiệp. Agent này nhận "
        "kết quả từ Search Agent (đã qua Judge) và Search Web Agent, sau đó tổng hợp, phân "
        "tích và đưa ra báo cáo có cấu trúc gồm 4 trường:")
    bullet(doc, "real_estate_findings: tóm tắt dạng markdown danh sách bài đăng + thông tin từ Internet, có kèm link nguồn.")
    bullet(doc, "summary_real_estate_findings: tóm tắt ngắn (không có link), tối đa 2000 ký tự.")
    bullet(doc, "analytics_and_advice: phân tích chi tiết + lời khuyên cá nhân hóa theo câu hỏi.")
    bullet(doc, "follow_up_questions: 3 câu hỏi tiếp theo có thể quan tâm.")
    para(doc,
        "Writer Agent dùng OpenAI Structured Output (response_format=Pydantic schema) để đảm "
        "bảo output luôn parse được. Prompt có RULES nghiêm ngặt cấm bịa số liệu — phải dùng "
        "EXACT giá / diện tích / địa chỉ từ data, ghi \"Chưa có thông tin\" cho field rỗng "
        "thay vì điền bừa.")

    add_h3(doc, "3.3.7  Bộ nhớ cho các tác tử")
    para(doc,
        "Để đảm bảo tính liền mạch và tự nhiên trong các cuộc hội thoại, mỗi tác tử trong "
        "hệ thống sử dụng kiến trúc bộ nhớ phân cấp:")
    bold_label(doc, "Bộ nhớ ngắn hạn (Short-term — Query Rewriter):",
        "Lưu 16 tin nhắn gần nhất (~8 lượt hội thoại). Mỗi tin user giới hạn 800 ký tự, mỗi "
        "tin assistant giới hạn 2000 ký tự — đủ để chứa danh sách top 3-5 bài đăng. Rewriter "
        "dùng cửa sổ này để viết lại câu hỏi standalone.")
    bold_label(doc, "Bộ nhớ dài hạn (Long-term — Zep Cloud):",
        "Triển khai dưới dạng đồ thị tri thức (Knowledge Graph) với Zep AI. Mỗi cuộc chat "
        "có Zep user riêng (user_id = \"chat-{chat_id}\") — bộ nhớ giữa các cuộc chat khác "
        "nhau được CÁCH LY hoàn toàn, không bị lẫn fact giữa cuộc này sang cuộc khác. Mỗi "
        "turn hoàn tất sẽ async lưu (user_msg, assistant_msg) vào Zep thread; turn sau gọi "
        "thread.get_user_context() để lấy summary làm bối cảnh.")
    para(doc,
        "Hai loại bộ nhớ bổ sung nhau: Rewriter cho precision (resolve refer cụ thể), Zep "
        "cho recall dài hạn (50+ turn nén thành knowledge graph). Đặc biệt việc isolate "
        "per chat trong Zep là cải tiến quan trọng của phiên bản cập nhật — bản gốc dùng "
        "chung 1 user_id mặc định nên các chat khác nhau bị share knowledge graph, dẫn đến "
        "context lẫn lộn.")

    add_h3(doc, "3.3.8  Truy xuất thông tin và phương pháp tìm kiếm")
    para(doc, "Để tối ưu hóa quá trình truy xuất thông tin, hệ thống sử dụng ba phương pháp:")
    bullet(doc, "Cosine Similarity: xác định độ tương đồng giữa các thực thể dựa trên ngữ nghĩa.")
    bullet(doc, "Okapi BM25 (full-text): đo độ tương đồng giữa các đoạn văn bản dựa trên từ khóa, dùng cho title + description trong Elasticsearch.")
    bullet(doc, "Breadth-First Search (BFS): trên knowledge graph Zep, mở rộng tìm kiếm từ các nút lân cận.")
    para(doc, "Kết quả tìm kiếm được sắp xếp lại (rerank) theo thời gian và mức độ liên quan, đảm bảo tính chính xác và phù hợp của câu trả lời.")

    add_h2(doc, "3.4  Đề xuất giải pháp Xây dựng Web Chatbot và Trực quan hóa Dữ liệu")
    para(doc,
        "Hệ thống tư vấn bất động sản sử dụng mô hình chatbot và trực quan hóa dữ liệu, "
        "được xây dựng với mục tiêu cung cấp cho người dùng một giao diện dễ sử dụng để tra "
        "cứu thông tin bất động sản, trò chuyện với chatbot và theo dõi các dữ liệu thống "
        "kê. Phiên bản cập nhật đã tinh gọn lại kiến trúc frontend: bỏ phần Đăng ký/Đăng "
        "nhập (MongoDB user auth) — không phù hợp với mô hình demo bài tập lớn — và thay "
        "bằng cơ chế multi-chat persist ở phía frontend qua file JSON.")
    add_h3(doc, "3.4.1  Cấu trúc hệ thống")
    bullet(doc, "Cơ sở dữ liệu bất động sản: Elasticsearch — lưu trữ 6 index tương ứng 6 loại BĐS, hỗ trợ full-text + aggregation.")
    bullet(doc, "Backend: FastAPI — endpoint /chat/, /chat/name_conversation, /dashboard/* và các legacy /get_price_*.")
    bullet(doc, "Frontend: Streamlit với theme Catppuccin Mocha tùy chỉnh — 2 tab chính Chatbot + Dashboard.")
    bullet(doc, "Bộ nhớ tác tử: Zep Cloud (long-term) + Query Rewriter (short-term).")
    bullet(doc, "Observability: Langfuse — trace mọi LLM call qua OpenTelemetry, debug + cost tracking.")

    add_h3(doc, "3.4.2  Quy trình hoạt động của hệ thống")
    bold_label(doc, "a. Hỏi đáp với Chatbot:", "")
    numbered(doc, "Người dùng vào tab Chatbot, nếu là cuộc chat trống sẽ thấy welcome screen với 4 sample prompt mẫu — click để gửi luôn.")
    numbered(doc, "Người dùng gõ câu hỏi → frontend gửi POST /chat/ kèm chat_id để backend xác định Zep thread.")
    numbered(doc, "Backend: load memory_context từ Zep → build payload [system + history + user] → ResearchManager.run() chạy pipeline 6 agent → trả về JSON (real_estate_findings, analytics_and_advice, follow_up_questions).")
    numbered(doc, "Frontend hiển thị message với spinner \"Đang suy nghĩ...\" trong lúc backend xử lý, scroll user message lên đầu viewport (pattern Claude/ChatGPT), tự động save state xuống file JSON cạnh module.")
    bold_label(doc, "b. Trực quan hóa dữ liệu:", "")
    numbered(doc, "Người dùng chọn tab Dashboard, filter tỉnh + loại BĐS.")
    numbered(doc, "Backend gọi /dashboard/kpi, /dashboard/price_segments, /dashboard/field_dist… → trả aggregate từ Elasticsearch.")
    numbered(doc, "Frontend hiển thị 3 KPI cards + Bar quận + Phân khúc giá + Pie đặc điểm (per estate type, không hiển thị field coverage thấp) + Heatmap giá trên bản đồ Folium.")
    page_break(doc)


def chapter_4(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 4. TRIỂN KHAI HỆ THỐNG")
    doc.add_paragraph()

    add_h2(doc, "4.1  Môi trường và chiến lược triển khai hệ thống")
    para(doc,
        "Hệ thống dữ liệu của project được thiết kế theo kiến trúc phân tán, bao gồm nhiều "
        "thành phần đảm nhiệm các chức năng khác nhau như thu thập dữ liệu, xử lý dữ liệu, "
        "lưu trữ, phân tích và trực quan hóa. Các thành phần lưu trữ và xử lý dữ liệu được "
        "triển khai theo mô hình container hóa và điều phối bởi Kubernetes (K8s) trên AWS "
        "EC2 nhằm đảm bảo tính linh hoạt, khả năng mở rộng và dễ dàng quản lý. Thành phần "
        "thu thập dữ liệu được triển khai local để tránh chi phí EC2 cho task định kỳ.")
    para(doc,
        "Trong giai đoạn phát triển sau bảo vệ, nhóm đã bổ sung pipeline dump/restore dữ "
        "liệu từ cluster ES trên AWS về Docker Compose local (`local_es/`) để dev và demo "
        "không phụ thuộc kết nối mạng AWS. Quy trình chi tiết được mô tả trong "
        "CAP_NHAT_DATA.md.")
    add_h3(doc, "4.1.1  Môi trường triển khai phần cứng và hệ điều hành")
    bullet(doc, "Máy cá nhân (developer): Windows 11 + WSL Ubuntu 20.04, Intel i7-1165G7 4 nhân, 8GB RAM. Dùng cho crawler, Airflow scheduler, frontend Streamlit, backend FastAPI.")
    bullet(doc, "4 instance EC2 m7i-flex (Ubuntu 20.04, x86_64, 4GiB RAM/1 vCPU): triển khai Kafka cluster, MinIO data lake, Elasticsearch cluster (K8s), Spark Streaming/Batch.")
    bullet(doc, "Local Docker Compose (`local_es/`): mirror ES + Kibana 8.11.1, dùng cho dev không đụng AWS.")

    add_h2(doc, "4.2  Triển khai thành phần thu thập và tiền xử lý dữ liệu")
    para(doc, "Thành phần thu thập và tiền xử lý được triển khai gồm hai dịch vụ chính:")
    bullet(doc, "Dịch vụ thu thập dữ liệu: trích xuất từ 2 trang batdongsan68.com.vn và bds68.com.vn qua Scrapy Spider.")
    bullet(doc, "Dịch vụ lập lịch: Apache Airflow chạy DAG theo cron hàng giờ.")
    add_h3(doc, "4.2.1  Dịch vụ thu thập dữ liệu")
    para(doc, "Mỗi loại BĐS (nhà phố, nhà riêng, chung cư, biệt thự, đất) có Spider riêng với schema Items đầy đủ 10+ trường. Crawler gọi qua FastAPI POST endpoint với 3 tham số chính: min_page, max_page, estate_type.")
    add_h3(doc, "4.2.2  Tiền xử lý dữ liệu")
    para(doc, "Hai bước: hiệu chỉnh địa chỉ (strip prefix, fuzzy match district/ward qua geoapivietnam khi không khớp) và loại bỏ trùng lặp (TF-IDF + cosine similarity title >0.95 hoặc description >0.99, danh sách so sánh giới hạn 4000 bài FIFO 75%).")
    add_h3(doc, "4.2.3  Xử lý dữ liệu thời gian thực")
    para(doc, "Sau tiền xử lý, dữ liệu được đẩy vào Kafka topic tương ứng (6 topic: nhamatpho, nharieng, chungcu, bietthu, dat, khac). Kafka cluster trên K8s với 3 broker để chịu lỗi, expose qua NodePort 32047.")
    add_h3(doc, "4.2.4  API thu thập dữ liệu")
    para(doc, "Cung cấp dưới dạng FastAPI POST với 3 tham số (min_page/max_page/estate_type). Airflow gọi periodically với chiến lược: chung cư/nhà riêng 100 bài/giờ; biệt thự/nhà phố 60 bài/giờ.")
    add_h3(doc, "4.2.5  Dịch vụ lập lịch thu thập dữ liệu")
    para(doc, "Apache Airflow đóng gói trong docker-compose, hoạt động trên cùng mạng với dịch vụ thu thập. DAG đơn giản: HttpOperator gọi API crawler theo cron @hourly.")

    add_h2(doc, "4.3  Triển khai thành phần lưu trữ và xử lý dữ liệu")
    add_h3(doc, "4.3.1  Tổng quan kiến trúc")
    para(doc, "Phân chia 3 layer: Staging (Kafka 3 brokers), Data Lake (MinIO Parquet) và Serving (Elasticsearch 6 index).")
    add_h3(doc, "4.3.2  Lưu trữ Dữ liệu")
    para(doc, "MinIO làm data lake lưu raw JSON + processed Parquet, hỗ trợ replay khi đổi logic. Elasticsearch lưu 6 index flat schema: nhamatpho_index, nharieng_index, chungcu_index, bietthu_index, dat_index, khac_index. Dữ liệu thực tế (snapshot 2026-05-21): ~87.000 bản ghi sống, phân bố như Bảng 4.1.")

    # Bảng 4.1: Coverage matrix per estate type
    add_caption(doc, "Bảng 4.1: Số lượng bản ghi và coverage % field theo 5 loại BĐS")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Field", "Nhà phố", "Nhà riêng", "Chung cư", "Biệt thự", "Đất"]):
        hdr[i].text = h
    rows = [
        ("Tổng bản ghi", "13,764", "15,550", "17,499", "12,252", "24,745"),
        ("square (diện tích)", "96.4%", "96.9%", "97.0%", "98.4%", "100%"),
        ("price", "90.9%", "91.6%", "85.4%", "87.0%", "90.4%"),
        ("price/square", "89.9%", "90.0%", "83.6%", "86.0%", "90.4%"),
        ("no_bedrooms", "67.5%", "80.0%", "85.3%", "64.8%", "1.2%"),
        ("no_bathrooms", "64.0%", "75.4%", "81.8%", "60.6%", "1.0%"),
        ("no_floors", "72.9%", "74.3%", "25.6%", "75.5%", "0.8%"),
        ("front_face (mặt tiền)", "31.6%", "31.8%", "0.4%", "40.3%", "73.0%"),
        ("front_road (lộ giới)", "25.5%", "29.6%", "0.5%", "36.1%", "70.5%"),
        ("direction (hướng)", "6.5%", "7.6%", "7.5%", "10.1%", "18.1%"),
    ]
    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = v

    para(doc, "")
    add_h3(doc, "4.3.3  Xử lý dữ liệu")
    para(doc, "Spark Structured Streaming đọc 6 Kafka topic, biến đổi dữ liệu qua 5 bước: định danh _id (post_id hoặc MD5(link)); chuẩn hóa text (loại unicode lạ); trích xuất feature số (chuẩn hóa giá VND, tính price/square); flatten extra_infos (no_bedrooms/no_floors/direction…); ghi song song MinIO + Elasticsearch. Batch processing qua Airflow gọi Spark định kỳ cho re-processing, aggregation, consistency check.")

    add_h2(doc, "4.4  Triển khai hệ thống đa tác tử")
    para(doc,
        "Phiên bản cập nhật chuyển stack chính thức từ Google Generative AI (Gemini) sang "
        "OpenAI Python SDK. Lý do: Gemini free tier có quota 0 RPM cho model pro, gây fail "
        "trong demo; OpenAI gpt-4o-mini có giá hợp lý ($0.15/1M input tokens) và structured "
        "output ổn định hơn. Tất cả 6 agent cùng dùng AsyncOpenAI/OpenAI client.")
    bold_label(doc, "Stack model thực tế:", "")
    bullet(doc, "Query Rewriter, Planner, Search DB, Search Web, Judge fallback: gpt-4o-mini (rẻ, nhanh, đủ chính xác).")
    bullet(doc, "Judge chính, Writer: gpt-4o (cần reasoning sâu cho đánh giá kết quả và viết báo cáo có cấu trúc).")
    bold_label(doc, "Search Database Agent — Cấu hình tools:",
        "Khai báo JSON schema thủ công cho 2 function (search_posts, search_posts_strict). "
        "Schema mirror đúng signature của hàm Python trong elasticsearch_queries.py, gồm 19 "
        "trường: estate_type[], province[], district[], ward[], front_face, front_road, "
        "no_bathrooms, no_bedrooms, no_floors, ultilization_square, price, min_price, "
        "max_price, price_per_square, square, description, is_latest_posted, "
        "is_latest_created. Việc lựa chọn được điều phối qua prompt với 8 rule chi tiết "
        "(xử lý đơn vị giá VND, đơn giá m², địa điểm có prefix Quận/Huyện, ưu tiên strict "
        "khi ≥2 ràng buộc).")
    bold_label(doc, "Sanitize input từ LLM:",
        "Class SearchAgent có method _sanitize_args() để chuẩn hóa tham số trả về từ "
        "OpenAI: ép int cho số phòng (LLM hay trả 2.0 thay vì 2), wrap string thành list "
        "cho fields tỉnh/quận, strip prefix \"Quận \"/\"Huyện \" (ES .keyword không match "
        "khi có prefix), ép float cho price/square/front_face.")
    bold_label(doc, "Judge Agent — Self-correction loop:",
        "Tối đa 2 vòng. Nếu vòng 1 needs_improvement, feedback sẽ được concat vào query "
        "gốc thành: \"Yêu cầu gốc: {query}. LƯU Ý ĐIỀU CHỈNH: {feedback}\" và truyền lại "
        "Search Agent. Nếu vòng 2 vẫn fail, accept kết quả hiện tại để không deadlock.")
    bold_label(doc, "Query Rewriter (mới):",
        "File custom_agents/query_rewriter.py. Dùng OpenAI structured output với schema "
        "RewrittenQuery(standalone_query: str). Truncate user 800 chars, assistant 2000 "
        "chars (đủ chứa top 3-5 listings). Temperature 0.1 để rewrite ổn định, gần "
        "deterministic. Khi history rỗng hoặc câu hỏi đã standalone, trả về nguyên văn — "
        "tiết kiệm 1 LLM call.")
    bold_label(doc, "Manager class (orchestrator):",
        "Class ResearchManager (manager.py) đóng vai trò orchestrator, KHÔNG phải LLM "
        "agent. Hàm run() tổ chức flow: rewrite → plan → (search_db loop judge | "
        "search_web) → write → async save Zep. Đây là điểm khác biệt quan trọng so với "
        "bản gốc — bản gốc gọi luôn class này là \"Manager Agent\", gây nhầm lẫn với role "
        "chọn tool (giờ tách thành Planner Agent).")

    add_h2(doc, "4.5  Triển khai thành phần Backend")
    para(doc,
        "Backend triển khai bằng FastAPI, kết nối với Elasticsearch (ES), OpenAI API "
        "(thay Gemini), Zep Cloud (memory) và Langfuse (observability). Phiên bản cập nhật "
        "bỏ MongoDB user auth — không có đăng ký/đăng nhập trong demo bài tập lớn — chat "
        "persist hoàn toàn ở phía frontend qua file JSON `.chat_state.json` cạnh app.py.")
    add_h3(doc, "4.5.1  Tương tác hỏi đáp với Agents")
    bullet(doc, "POST /chat/: pipeline chính. Body là list[messages] gửi từ frontend (item cuối phải kèm chat_id). Trả {real_estate_findings, analytics_and_advice, follow_up_questions}.")
    bullet(doc, "POST /chat/name_conversation: gọi sau turn đầu tiên để tự sinh tên cuộc hội thoại dạng \"Mua nhà Hoàn Kiếm 50 tỷ\" (~3-8 từ).")
    add_h3(doc, "4.5.2  Endpoints Dashboard v2")
    para(doc, "8 endpoint /dashboard/* trả aggregate từ Elasticsearch cho frontend:")
    bullet(doc, "GET /dashboard/kpi: total_listings + median_price + median_price_per_sq + new_7d.")
    bullet(doc, "GET /dashboard/price_segments: phân khúc giá <2 tỷ, 2-5, 5-10, 10-20, >20 tỷ.")
    bullet(doc, "GET /dashboard/pps_quartiles: quartile giá/m² theo quận.")
    bullet(doc, "GET /dashboard/field_dist/{field}: distribution categorical (no_bedrooms, no_floors).")
    bullet(doc, "GET /dashboard/range_dist/{field}: distribution range (square, front_face).")
    bullet(doc, "GET /dashboard/trend_monthly: trend median theo tháng (cho roadmap tab Nhà đầu tư).")
    bullet(doc, "Cộng các legacy /get_price_* giữ lại cho dashboard v1 backward compat.")
    add_h3(doc, "4.5.3  Set active province")
    para(doc, "POST /set_active_province/{province}: cập nhật biến global valid_districts theo tỉnh user chọn (đọc từ province_district_ward_prefix.json), mọi query dashboard sau đó filter theo list này.")
    add_h3(doc, "4.5.4  Các thành phần hỗ trợ")
    bullet(doc, "Zep Cloud SDK 3.x: thread.get_user_context (cross-thread knowledge graph) + thread.add_messages (lưu turn async).")
    bullet(doc, "Langfuse OTLP: trace mọi span (Rewriter, Planner, SearchDB, Judge, Writer) qua OpenTelemetry. Fail silent nếu thiếu LANGFUSE_PUBLIC_KEY/SECRET_KEY.")
    bullet(doc, "Rich Console Printer: pretty-print log từng bước cho debug khi chạy uvicorn.")

    add_h2(doc, "4.6  Triển khai thành phần trực quan hóa và giao diện Chatbot")
    para(doc,
        "Frontend phát triển bằng Streamlit, theme tùy chỉnh dựa trên Catppuccin Mocha "
        "palette (nền base #1e1e2e, accent mauve #cba6f7, blue #89b4fa) — đồng bộ với "
        "nvim/wezterm setup của nhóm. Bỏ hoàn toàn UI Đăng ký/Đăng nhập của bản cũ. Cấu "
        "trúc chính: 1 hero header gradient + 2 tab top-nav (Chatbot / Dashboard).")
    add_h3(doc, "4.6.1  Giao diện Chatbot")
    bullet(doc, "Layout 2 cột: chat list panel bên trái (fixed via position: fixed override Streamlit overflow), chat area bên phải.")
    bullet(doc, "Multi-chat persist: state lưu file .chat_state.json cạnh app.py (gồm active_chat_id + dict chats), reload tab vẫn giữ nguyên history. Tự migrate format cũ (single-chat) sang format mới (multi-chat) nếu phát hiện.")
    bullet(doc, "Empty state khi cuộc chat trống: welcome card có icon, heading gradient + 4 sample prompt mẫu (\"Mua nhà mặt phố Hoàn Kiếm 50 tỷ\", \"Chung cư 2PN Cầu Giấy 3-5 tỷ\", \"Đất nền Đà Nẵng dưới 5 tỷ\", \"Biệt thự liền kề Hoài Đức\") — click button tự động gửi luôn.")
    bullet(doc, "Auto-scroll: khi user gửi message, JS scroll-into-view block:'start' đẩy user message lên ngay dưới tab bar (pattern Claude/ChatGPT), kèm spinner \"🤔 Đang suy nghĩ...\" trong bubble assistant cho tới khi response về.")
    bullet(doc, "Chat input fixed bottom + tab nav fixed top (CSS position: fixed + scroll-margin-top bù chiều cao tab bar).")
    add_h3(doc, "4.6.2  Dashboard")
    para(doc,
        "Filter trên cùng: selectbox 63 tỉnh thành + radio 5 loại BĐS (nhà mặt phố / nhà "
        "riêng / chung cư / biệt thự / đất). Các widget được trim so với bản gốc theo "
        "coverage data thật:")
    bullet(doc, "3 KPI cards: Tổng tin đăng / Giá điển hình (median) / Giá/m² điển hình. (Bản gốc có 4 thẻ — đã bỏ \"Tin mới 7 ngày\" vì coverage thấp.)")
    bullet(doc, "Bar quận/huyện: Giá Trung Bình (peach) + Đơn giá /m² (sapphire).")
    bullet(doc, "Bar Phân phối giá theo phân khúc bins.")
    bullet(doc, "Pie Đặc điểm phổ biến — cấu hình per estate type (Bảng 4.2). Slice \"Chưa rõ\" tô màu xám overlay0 để honest về missing data.")
    bullet(doc, "Heatmap giá trung bình trên bản đồ Folium per tỉnh (load coord từ district_coords_full.json).")

    # Bảng 4.2: DISTRIBUTION_LAYOUT
    add_caption(doc, "Bảng 4.2: Cấu hình widget Pie Đặc điểm theo loại BĐS")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Loại BĐS"
    table.rows[0].cells[1].text = "Widget hiển thị"
    layout_rows = [
        ("Nhà mặt phố (nhamatpho)", "Số tầng"),
        ("Nhà riêng (nharieng)", "Số phòng ngủ + Số tầng"),
        ("Chung cư (chungcu)", "Số phòng ngủ + Số phòng tắm"),
        ("Biệt thự (bietthu)", "Số phòng ngủ + Số tầng"),
        ("Đất (dat)", "Phân khúc diện tích + Mặt tiền"),
        ("Khác (khac)", "Phân khúc diện tích + Số tầng + Mặt tiền"),
    ]
    for r in layout_rows:
        cells = table.add_row().cells
        cells[0].text = r[0]
        cells[1].text = r[1]
    page_break(doc)


def chapter_5(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 5. ĐÁNH GIÁ THỰC NGHIỆM")
    para(doc, "[TODO: nội dung Chương 5]", indent_first=False)
    page_break(doc)


def chapter_6(doc: Document) -> None:
    add_h1(doc, "CHƯƠNG 6. KẾT LUẬN")
    para(doc, "[TODO: nội dung Chương 6]", indent_first=False)
    page_break(doc)


def references(doc: Document) -> None:
    add_h1(doc, "TÀI LIỆU THAM KHẢO")
    para(doc, "[TODO: refs]", indent_first=False)


# ============================================================
# MAIN
# ============================================================

def main() -> None:
    doc = Document()
    setup_styles(doc)
    page_setup(doc)

    build_cover(doc)
    build_foreword(doc)
    build_toc_placeholder(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    references(doc)

    doc.save(OUTPUT)
    print(f"✅ Saved: {OUTPUT}")
    print(f"   Word: mở file → Ctrl+A → F9 để cập nhật Mục lục.")


if __name__ == "__main__":
    main()
